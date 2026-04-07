import re
import sys
from pathlib import Path


ROOT = Path(__file__).resolve().parent
PYDEPS = ROOT / ".pydeps"
if PYDEPS.exists():
    sys.path.insert(0, str(PYDEPS))

from docx import Document  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Pt  # type: ignore


MD_PATH = ROOT / "本系统毕业论文初稿.md"
TEMPLATE_PATH = ROOT / "论文格式内容模板.docx"
OUT_PATH = ROOT / "本系统毕业论文初稿_修订版_v8.docx"


def _set_run_font(run, cn="宋体", en="Times New Roman", size_pt=12):
    run.font.name = en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
    run.font.size = Pt(size_pt)


def _clean_text(text: str) -> str:
    text = text.replace("\\*", "*")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("*", "")
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def _add_text_with_superscript(paragraph, text: str):
    cleaned = _clean_text(text)
    if not cleaned:
        return
    last = 0
    for m in re.finditer(r"<sup>(.*?)</sup>", cleaned, flags=re.IGNORECASE):
        normal = cleaned[last : m.start()]
        if normal:
            run = paragraph.add_run(normal)
            _set_run_font(run)
        sup = re.sub(r"\s+", "", m.group(1))
        if sup:
            run = paragraph.add_run(sup)
            _set_run_font(run, size_pt=10.5)
            run.font.superscript = True
        last = m.end()
    tail = cleaned[last:]
    if tail:
        run = paragraph.add_run(tail)
        _set_run_font(run)


def _clear_document(doc: Document):
    body = doc._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    _add_text_with_superscript(p, text)
    return p


def _add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(_clean_text(line.rstrip("\n")))
        _set_run_font(run, cn="等线", en="Consolas", size_pt=10.5)


def _parse_table(block: list[str]):
    def split_row(row: str):
        row = row.strip()
        if row.startswith("|"):
            row = row[1:]
        if row.endswith("|"):
            row = row[:-1]
        return [_clean_text(c.strip()) for c in row.split("|")]

    header = split_row(block[0])
    rows = [split_row(r) for r in block[2:]]
    col_count = max(len(header), *(len(r) for r in rows)) if rows else len(header)
    header = header + [""] * (col_count - len(header))
    rows = [r + [""] * (col_count - len(r)) for r in rows]
    return header, rows


def _add_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    for j, val in enumerate(header):
        cell = table.cell(0, j)
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                _set_run_font(r)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_run_font(r)


def _is_hr(line: str) -> bool:
    s = line.strip()
    return s in {"---", "***", "___"}


def _is_heading(line: str):
    m = re.match(r"^(#{1,6})\s+(.*)$", line.rstrip())
    if not m:
        return None
    level = len(m.group(1))
    text = m.group(2).strip()
    return level, text


def _is_table_header(line: str, next_line: str) -> bool:
    if "|" not in line:
        return False
    s = next_line.strip()
    if "|" not in s:
        return False
    s = s.strip("|").strip()
    parts = [p.strip() for p in s.split("|")]
    if not parts:
        return False
    return all(re.fullmatch(r"-{3,}:?", p) or re.fullmatch(r":?-{3,}:?", p) for p in parts)


def _is_image_line(line: str) -> bool:
    return bool(re.match(r"^\s*!\[.*?\]\(.*?\)\s*$", line))


def _is_reference_line(line: str) -> bool:
    return bool(re.match(r"^\s*\[\d+\]\s*", line))


def export_docx():
    if not MD_PATH.exists():
        raise FileNotFoundError(str(MD_PATH))

    md = MD_PATH.read_text(encoding="utf-8")
    lines = md.splitlines()

    if TEMPLATE_PATH.exists():
        doc = Document(str(TEMPLATE_PATH))
        _clear_document(doc)
    else:
        doc = Document()

    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    paragraph_buf: list[str] = []

    def flush_paragraph():
        nonlocal paragraph_buf
        text = "\n".join([t.rstrip() for t in paragraph_buf]).strip()
        paragraph_buf = []
        if not text:
            return
        _add_paragraph(doc, text)

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            fence = line.strip()
            if not in_code:
                flush_paragraph()
                in_code = True
                code_lang = fence.strip("`").strip()
                code_lines = []
            else:
                in_code = False
                if code_lang.lower().startswith("mermaid"):
                    _add_paragraph(doc, "图示代码（Mermaid）：")
                _add_code_block(doc, code_lines)
                code_lang = ""
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if _is_hr(line):
            flush_paragraph()
            doc.add_page_break()
            i += 1
            continue

        heading = _is_heading(line)
        if heading:
            flush_paragraph()
            level, text = heading
            p = doc.add_heading(_clean_text(text), level=level)
            for r in p.runs:
                _set_run_font(r)
            i += 1
            continue

        if i + 1 < len(lines) and _is_table_header(line, lines[i + 1]):
            flush_paragraph()
            table_block = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                table_block.append(lines[i])
                i += 1
            header, rows = _parse_table(table_block)
            _add_table(doc, header, rows)
            continue

        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        if _is_image_line(line):
            flush_paragraph()
            i += 1
            continue

        if _is_reference_line(line):
            flush_paragraph()
            _add_paragraph(doc, line)
            i += 1
            continue

        paragraph_buf.append(_clean_text(line))
        i += 1

    flush_paragraph()
    doc.save(str(OUT_PATH))


if __name__ == "__main__":
    export_docx()
