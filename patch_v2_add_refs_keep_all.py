import os
import sys
import time
from pathlib import Path

ROOT = Path(__file__).resolve().parent
PYDEPS = ROOT / ".pydeps"
if PYDEPS.exists():
    sys.path.insert(0, str(PYDEPS))

from docx import Document  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Pt  # type: ignore

DOCX_PATH = ROOT / "本系统毕业论文初稿_修订版_v2.docx"
TMP_PATH = ROOT / "本系统毕业论文初稿_修订版_v2.__tmp__.docx"
FALLBACK_PATH = ROOT / "本系统毕业论文初稿_修订版_v3.docx"

ADD_REFS = [
    "[1]李岸.大学生创新创业项目管理系统的设计与实现[D].广西大学,2021.",
    "[2]王珂.济南职业学院大学生创新创业管理系统的设计与实现[D].山东大学,2019.",
    "[3]陈婧.大学生创新创业项目管理系统的设计与实现[D].厦门大学,2018.",
]


def set_run_style(run, superscript=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    run.font.superscript = superscript


def add_abstract_citation(doc: Document):
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if "摘要" in t and "Abstract" not in t:
            for j in range(i + 1, min(i + 20, len(doc.paragraphs))):
                pp = doc.paragraphs[j]
                tt = pp.text.strip()
                if not tt:
                    continue
                if "关键词" in tt or "Abstract" in tt:
                    return False
                if "[1][2][3]" in tt:
                    return False
                run = pp.add_run("[1][2][3]")
                set_run_style(run, superscript=True)
                return True
    return False


def add_refs_keep_existing(doc: Document):
    full_text = "\n".join(p.text for p in doc.paragraphs)
    to_add = [r for r in ADD_REFS if r not in full_text]
    if not to_add:
        return 0

    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "参考文献" in p.text:
            ref_idx = i
            break

    if ref_idx is None:
        doc.add_paragraph("参考文献")
    for r in to_add:
        p = doc.add_paragraph()
        run = p.add_run(r)
        set_run_style(run, superscript=False)
    return len(to_add)


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(str(DOCX_PATH))
    doc = Document(str(DOCX_PATH))

    abs_added = add_abstract_citation(doc)
    refs_added = add_refs_keep_existing(doc)

    doc.save(str(TMP_PATH))
    try:
        os.replace(str(TMP_PATH), str(DOCX_PATH))
        print(f"saved={DOCX_PATH}, abstract_added={abs_added}, refs_added={refs_added}")
    except PermissionError:
        alt = FALLBACK_PATH
        if alt.exists():
            alt = ROOT / f"本系统毕业论文初稿_修订版_v4_{int(time.time())}.docx"
        doc.save(str(alt))
        try:
            if TMP_PATH.exists():
                TMP_PATH.unlink()
        except Exception:
            pass
        print(f"saved={alt}, abstract_added={abs_added}, refs_added={refs_added}, note=v2_locked")


if __name__ == "__main__":
    main()
