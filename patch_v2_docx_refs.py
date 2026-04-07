import re
import sys
import os
from pathlib import Path

ROOT = Path(__file__).resolve().parent
PYDEPS = ROOT / ".pydeps"
if PYDEPS.exists():
    sys.path.insert(0, str(PYDEPS))

from docx import Document  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Pt  # type: ignore

DOCX_PATH = ROOT / "本系统毕业论文初稿_修订版_v2.docx"
TMP_PATH = ROOT / "本系统毕业论文初稿_修订版_v2.__tmp__.docx"
FALLBACK_PATH = ROOT / "本系统毕业论文初稿_修订版_v3.docx"

NEW_REFS = {
    1: "[1]李岸.大学生创新创业项目管理系统的设计与实现[D].广西大学,2021.",
    2: "[2]王珂.济南职业学院大学生创新创业管理系统的设计与实现[D].山东大学,2019.",
    3: "[3]陈婧.大学生创新创业项目管理系统的设计与实现[D].厦门大学,2018.",
}


def set_run_style(run, superscript=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    run.font.superscript = superscript


def clear_paragraph(paragraph):
    for r in list(paragraph.runs):
        paragraph._element.remove(r._element)


def replace_reference_1_3(doc: Document):
    start = None
    for idx, p in enumerate(doc.paragraphs):
        if "参考文献" in p.text:
            start = idx
            break
    if start is None:
        return 0

    changed = 0
    existing_nums = set()
    pat = re.compile(r"^\s*\[(\d+)\]")
    for p in doc.paragraphs[start + 1 :]:
        t = p.text.strip()
        m = pat.match(t)
        if not m:
            continue
        n = int(m.group(1))
        existing_nums.add(n)
        if n in NEW_REFS:
            clear_paragraph(p)
            run = p.add_run(NEW_REFS[n])
            set_run_style(run, superscript=False)
            changed += 1
    for n in (1, 2, 3):
        if n not in existing_nums:
            p = doc.add_paragraph()
            run = p.add_run(NEW_REFS[n])
            set_run_style(run, superscript=False)
            changed += 1
    return changed


def add_abstract_citation(doc: Document):
    abs_idx = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "摘要" in text and "Abstract" not in text:
            abs_idx = i
            break
    if abs_idx is None:
        return False

    for j in range(abs_idx + 1, min(abs_idx + 25, len(doc.paragraphs))):
        p = doc.paragraphs[j]
        text = p.text.strip()
        if not text:
            continue
        if "Abstract" in text:
            return False
        if len(text) < 20:
            continue
        if "[1][2][3]" in text:
            return False
        run = p.add_run("[1][2][3]")
        set_run_style(run, superscript=True)
        return True
    return False


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(str(DOCX_PATH))

    doc = Document(str(DOCX_PATH))
    ref_count = replace_reference_1_3(doc)
    abs_changed = add_abstract_citation(doc)
    doc.save(str(TMP_PATH))
    try:
        os.replace(str(TMP_PATH), str(DOCX_PATH))
        print(f"saved={DOCX_PATH}, references_changed={ref_count}, abstract_citation_added={abs_changed}")
    except PermissionError:
        doc.save(str(FALLBACK_PATH))
        try:
            if TMP_PATH.exists():
                TMP_PATH.unlink()
        except Exception:
            pass
        print(f"saved={FALLBACK_PATH}, references_changed={ref_count}, abstract_citation_added={abs_changed}, note=v2_locked")


if __name__ == "__main__":
    main()
